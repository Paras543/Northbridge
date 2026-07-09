"""
app/services/extraction.py

Pulls plain text out of raw uploaded file bytes, based on content_type.
Pure function: bytes in, text out. No DB, no I/O — easy to unit test.
"""

import io
import csv

from pypdf import PdfReader
from docx import Document as DocxDocument


class ExtractionError(Exception):
    """Raised when text cannot be extracted from the given file."""


def extract_text(raw_bytes: bytes, content_type: str, filename: str = "") -> str:
    """
    Extract text from raw file bytes based on content_type.
    Falls back to filename extension if content_type is generic/missing.
    """
    content_type = (content_type or "").lower()
    filename_lower = (filename or "").lower()

    if "pdf" in content_type or filename_lower.endswith(".pdf"):
        return _extract_pdf(raw_bytes)

    if "wordprocessingml" in content_type or filename_lower.endswith(".docx"):
        return _extract_docx(raw_bytes)

    if "csv" in content_type or filename_lower.endswith(".csv"):
        return _extract_csv(raw_bytes)

    if content_type.startswith("text/") or filename_lower.endswith(".txt"):
        return _extract_txt(raw_bytes)

    # Old .doc (not .docx) and other binary formats aren't supported here.
    raise ExtractionError(
        f"Unsupported file type for extraction: content_type='{content_type}', filename='{filename}'"
    )


def _extract_pdf(raw_bytes: bytes) -> str:
    try:
        reader = PdfReader(io.BytesIO(raw_bytes))
    except Exception as e:
        raise ExtractionError(f"Could not read PDF: {e}") from e

    if reader.is_encrypted:
        raise ExtractionError("PDF is password-protected; cannot extract text.")

    pages_text = []
    for page in reader.pages:
        pages_text.append(page.extract_text() or "")

    text = "\n\n".join(pages_text).strip()
    if not text:
        
        raise ExtractionError(
            "No extractable text found in PDF (likely a scanned/image-only document)."
        )
    return text


def _extract_docx(raw_bytes: bytes) -> str:
    try:
        doc = DocxDocument(io.BytesIO(raw_bytes))
    except Exception as e:
        raise ExtractionError(f"Could not read DOCX: {e}") from e

    parts = [p.text for p in doc.paragraphs if p.text.strip()]

    
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells)
            if row_text.strip(" |"):
                parts.append(row_text)

    text = "\n".join(parts).strip()
    if not text:
        raise ExtractionError("No extractable text found in DOCX.")
    return text


def _extract_csv(raw_bytes: bytes) -> str:
    try:
        decoded = raw_bytes.decode("utf-8-sig")
    except UnicodeDecodeError:
        decoded = raw_bytes.decode("latin-1")

    reader = csv.reader(io.StringIO(decoded))
    lines = [", ".join(row) for row in reader]
    text = "\n".join(lines).strip()
    if not text:
        raise ExtractionError("CSV file appears to be empty.")
    return text


def _extract_txt(raw_bytes: bytes) -> str:
    try:
        text = raw_bytes.decode("utf-8")
    except UnicodeDecodeError:
        text = raw_bytes.decode("latin-1")

    text = text.strip()
    if not text:
        raise ExtractionError("Text file appears to be empty.")
    return text



